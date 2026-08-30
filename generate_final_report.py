# -*- coding: utf-8 -*-
"""Generate an Arabic academic DOCX report matching the Final Compiler Project style."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

OUT = r"C:\Users\user\Desktop\work_project\work_project\Final_Compiler_Project.docx"

AR_FONT = "Traditional Arabic"
EN_FONT = "Calibri"


def set_run_font(run, name=AR_FONT, size=14, bold=False, color=None, ltr=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = EN_FONT if ltr else name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    fname = EN_FONT if ltr else name
    rFonts.set(qn("w:ascii"), fname)
    rFonts.set(qn("w:hAnsi"), fname)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), fname)
    if not ltr:
        rtl = rPr.find(qn("w:rtl"))
        if rtl is None:
            rtl = OxmlElement("w:rtl")
            rPr.append(rtl)
        rtl.set(qn("w:val"), "1")
        cs = rPr.find(qn("w:cs"))
        if cs is None:
            cs = OxmlElement("w:cs")
            rPr.append(cs)
        cs.set(qn("w:val"), "1")
    if color:
        run.font.color.rgb = color


def set_paragraph_rtl(p, align="right"):
    pPr = p._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    mapping = {"right": "right", "center": "center", "left": "left", "both": "both"}
    jc.set(qn("w:val"), mapping.get(align, "right"))
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "both":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_ar(doc, text, size=14, bold=False, align="both", space_after=8, space_before=0):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading_ar(doc, text, size=18):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, "right")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))
    return p


def add_bullet(doc, text, size=14):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, "both")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run("❖  " + text)
    set_run_font(run, size=size)
    return p


def add_numbered(doc, n, text, size=14):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, "both")
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(f"{n}.  {text}")
    set_run_font(run, size=size)
    return p


def add_code(doc, code):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, "left")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.6)
    pf = p.paragraph_format
    shd = OxmlElement("w:shd")
    pPr = p._p.get_or_add_pPr()
    shd.set(qn("w:fill"), "F4F6F8")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    run = p.add_run(code)
    set_run_font(run, name=EN_FONT, size=11, ltr=True)
    run.font.name = "Consolas"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    return p


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.right_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    sectPr.append(bidi)
    add_page_number(section)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = AR_FONT
    normal.font.size = Pt(14)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), AR_FONT)
    rFonts.set(qn("w:hAnsi"), AR_FONT)
    rFonts.set(qn("w:cs"), AR_FONT)


def build():
    doc = Document()
    configure_doc(doc)

    # ── Cover ────────────────────────────────────────────────────────────
    for _ in range(2):
        doc.add_paragraph()

    add_ar(doc, "الجمهورية العربية السورية", size=16, bold=True, align="center", space_after=2)
    add_ar(doc, "جامعة دمشق — كلية الهندسة المعلوماتية", size=18, bold=True, align="center", space_after=4)
    add_ar(doc, "مقرر المترجمات", size=16, bold=True, align="center", space_after=18)

    add_ar(doc, "التقرير النهائي لمشروع المترجم", size=22, bold=True, align="center", space_after=6)
    add_ar(
        doc,
        "مترجم تطبيقات ويب ديناميكية باستخدام Flask و Jinja2 مع HTML و CSS",
        size=16,
        bold=True,
        align="center",
        space_after=20,
    )

    add_ar(doc, "العام الدراسي 2025 / 2026", size=14, bold=True, align="center", space_after=18)

    add_ar(doc, "إشراف المهندس العملي:", size=14, bold=True, align="center", space_after=2)
    add_ar(doc, "م. هيا شاهده", size=14, align="center", space_after=2)
    add_ar(doc, "م. سارة القبعة", size=14, align="center", space_after=18)

    add_ar(doc, "إعداد وتقديم الطلاب:", size=14, bold=True, align="center", space_after=6)
    add_ar(doc, "❖  ............................................", size=14, align="center", space_after=2)
    add_ar(doc, "❖  ............................................", size=14, align="center", space_after=2)
    add_ar(doc, "❖  ............................................", size=14, align="center", space_after=2)
    add_ar(doc, "❖  ............................................", size=14, align="center", space_after=2)
    add_ar(doc, "❖  ............................................", size=14, align="center", space_after=2)
    add_ar(doc, "❖  ............................................", size=14, align="center", space_after=12)

    doc.add_page_break()

    # ── أولاً ────────────────────────────────────────────────────────────
    add_heading_ar(doc, "أولاً: تعريف بلغة المشروع")

    add_ar(
        doc,
        "يهدف هذا المشروع إلى تطبيق المفاهيم النظرية للمترجمات تطبيقاً عملياً، عبر بناء مترجم "
        "يستهدف إطار العمل Flask ولغة القوالب Jinja2 مع HTML و CSS. المطلوب ليس كتابة تطبيق ويب "
        "يدوياً فحسب، بل بناء سلسلة ترجمة كاملة: تحليل لفظي ونحوي، ثم شجرتا تركيب مجردتان، ثم تحليل "
        "دلالي وجدول رموز، ثم توليد كود قادر على تشغيل واجهات ويب ديناميكية قابلة للتطوير.",
    )
    add_ar(
        doc,
        "اللغة التي بنيناها في المشروع هي مجموعة جزئية (subset) موجّهة لتطبيقات Flask العملية. "
        "قسّمنا البناء إلى جبهتين متكاملتين: واجهة Python/Flask من جهة، وواجهة قوالب HTML/Jinja2/CSS "
        "من جهة أخرى. يمكن استخدام اللغة الناتجة لتطوير تطبيقات ويب تعرض منتجات، تضيف منتجاً، تعرض "
        "تفاصيل منتج، وتحذف منتجاً، مع ضمان التنقل السلس بين هذه الواجهات. ومن أبرز ما تدعمه اللغة ما يلي:",
    )

    add_ar(doc, "أ) مجموعة Python / Flask", size=15, bold=True, align="right", space_before=8)
    add_bullet(doc, "الاستيراد باستخدام from و import، بما في ذلك استيراد رموز Flask مثل Flask و render_template و request و redirect و url_for.")
    add_bullet(doc, "تعريف التطبيق عبر Flask(__name__) وتعريف المسارات باستخدام المزخرف @app.route مع وسائط اختيارية مثل methods=[\"GET\", \"POST\"].")
    add_bullet(doc, "تعريف التوابع باستخدام def مع قائمة معاملات مستقلة، واستدعاؤها لاحقاً كمسارات أو كدوال عادية.")
    add_bullet(doc, "استخدام الجمل الشرطية if و elif و else بشكل متسلسل يسمح بتحقيق شروط مركّبة ببساطة.")
    add_bullet(doc, "استخدام الحلقة for والمتكررة while، مع عبارة break للتحكم بسير التعليمات داخل الحلقات.")
    add_bullet(doc, "استخدام return لإرجاع نتيجة واحدة أو أكثر (مثل رسالة مع رمز الحالة 404).")
    add_bullet(doc, "تعريف القيم بجميع أنواعها الممكنة في هذه المجموعة: None و number و string و boolean والقوائم arrays والكائنات dictionaries.")
    add_bullet(doc, "إسناد القيم للمتغيرات، والإسناد المركّب مثل += و -= و *= وغيرها.")
    add_bullet(doc, "الوصول إلى عناصر القوائم أو خصائص الكائنات عبر الفهرسة name[key] والنقطة name.attr.")
    add_bullet(doc, "استدعاء التوابع وإنشاء كائنات جديدة من صنف معيّن، واستدعاء التوابع كقيم يمكن إسنادها واستخدامها كـ callbacks.")
    add_bullet(doc, "استخدام عملية النفي المنطقي not والعمليات الحسابية من جمع وضرب وطرح وقسمة وباقي وقوة، مع الأولويات الرياضية الصحيحة.")
    add_bullet(doc, "استخدام عمليات المقارنة المنطقية (أكبر، أكبر أو يساوي، أصغر، أصغر أو يساوي) والتحقق من المساواة أو عدم المساواة بين طرفين، إضافة إلى is و is not.")
    add_bullet(doc, "استخدام العمليات المنطقية And و Or.")
    add_bullet(doc, "استخدام المعامل الثلاثي (العبارة الشرطية القصيرة أو Ternary Operator) بالشكل x if cond else y.")
    add_bullet(doc, "دعم قوائم الإدراك List Comprehension بالشكل [expr for x in y if cond].")
    add_bullet(doc, "استخدام كائنات وتوابع من أصل اللغة (built-in) لا يحتاج المستخدم إلى تعريفها بنفسه، مثل print و float و int و str و len و max و min، إضافة إلى رموز Flask الجاهزة في النطاق العام.")

    add_ar(doc, "ب) مجموعة Jinja2 / HTML / CSS", size=15, bold=True, align="right", space_before=8)
    add_bullet(doc, "استخدام صياغة HTML مباشرة داخل القالب: عناصر عادية، وعناصر ذات إغلاق ذاتي، وعناصر فارغة (void) مثل img و br و hr و meta و link و input.")
    add_bullet(doc, "تخصيص خصائص المعلومات لكل عنصر عن طريق Attributes، ويمكن أن تكون قيمة الخاصية نصاً ثابتاً أو تعبيراً Jinja.")
    add_bullet(doc, "استخدام التعابير {{ variable }} و {{ product.field }} لاستبدال المتغيرات بقيم آتية من بيانات Python.")
    add_bullet(doc, "استخدام الحلقة {% for x in y %} ... {% endfor %} لتكرار جزء من الواجهة على قائمة مستخرجة من برنامج Python.")
    add_bullet(doc, "دعم {% extends %} و {% block %} و {% endblock %} على مستوى التحليل النحوي والدلالي.")
    add_bullet(doc, "استدعاء توابع مدمجة مثل url_for داخل القالب لتوليد مسارات التنقل والملفات الثابتة.")
    add_bullet(doc, "تحليل CSS داخل وسم <style>: محددات العنصر والصنف والمعرّف والمحدد العام، مع تصريحات الخاصية والقيمة.")

    add_ar(
        doc,
        "ما سبق إضافة إلى تفاصيل أخرى حققناها في هذا المشروع، وسنتعرّف على كيفية قيامنا بذلك في الأقسام التالية من التقرير.",
        space_before=8,
    )

    # ── ثانياً ───────────────────────────────────────────────────────────
    add_heading_ar(doc, "ثانياً: بناء قواعد اللغة اللفظية والنحوية")

    add_ar(
        doc,
        "في البداية قمنا بكتابة قواعد اللغة اللفظية والنحوية باستخدام لغة قواعد ANTLR (الملفات ذات اللاحقة g4). "
        "وبموجب إعلان المشروع وجب تعريف الصياغة الخاصة بـ Python و Jinja2 و HTML و CSS، لذلك أنشأنا مجموعتين من الملفات:",
    )
    add_bullet(doc, "ملف pyLexer.g4 الذي يحوي القواعد اللفظية لمجموعة Python/Flask، وملف pyParser.g4 الذي يحوي القواعد النحوية.")
    add_bullet(doc, "ملف htmlLexer.g4 الذي يحوي القواعد اللفظية لـ HTML و Jinja2 و CSS عبر أنماط lexer متعددة (modes)، وملف htmlParser.g4 الذي يحوي القواعد النحوية للقوالب.")

    add_ar(
        doc,
        "حيث تم تجميع رموز لغة Python وكتابتها في الملف الخاص بالـ Lexer على شكل Tokens، ثم إنشاء القواعد وكتابتها في الملف الخاص بالـ Parser. "
        "ومن أهم قرارات التصميم في محلّل Python: وضع الكلمات المفتاحية قبل المعرّفات حتى لا تُبتلع كأسماء، ووضع الرموز الأطول قبل الأقصر (مثل **= قبل **)، "
        "ومعالجة المسافات البادئة عبر توكنَي INDENT و DEDENT بمساعدة الصنف PyLexerBase. أما التعبيرات فقد بُنيت كبرج أولويات من الأدنى إلى الأعلى: "
        "الثلاثي ثم or ثم and ثم not ثم المقارنة ثم الجمع فالضرب فالقوة فالأحادي ثم اللاحقة (استدعاء، فهرسة، خاصية) ثم الذرة.",
    )
    add_ar(
        doc,
        "أما محلّل القوالب فيعتمد أنماطاً متداخلة في الـ Lexer: النمط الافتراضي للنص ووسم البداية، ونمط TAG لوسوم HTML، ونمط JINJA_STMT للجمل {% ... %}، "
        "ونمط JINJA_EXPR للتعابير {{ ... }}، ونمط CSS لمحتوى <style>. هذا الفصل يمنع تداخل رموز اللغات الأربع ويُبقي القواعد النحوية واضحة: "
        "القالب سلسلة محتوى، والمحتوى قد يكون doctype أو عنصراً HTML أو عنصراً style أو جملة Jinja أو تعبيراً Jinja أو نصاً.",
    )
    add_ar(
        doc,
        "وبعد ذلك قمنا باستخدام أداة ANTLR لتوليد الـ Lexer (المحلّل اللفظي) والـ Parser (المحلّل القواعدي) بناءً على هذه القواعد، "
        "وبعض الأصناف الأخرى التي تفيدنا في بناء المترجم الخاص بنا مثل أصناف العقد التي سنبني منها شجرة AST، "
        "وكذلك الأصناف الخاصة بالتقاط ومعالجة الأخطاء وأصناف جدول الرموز وغير ذلك من الأصناف الضرورية المساعدة والمكمّلة.",
    )

    # ── ثالثاً ───────────────────────────────────────────────────────────
    add_heading_ar(doc, "ثالثاً: توليد شجرة AST")

    add_ar(
        doc,
        "الأصناف التي تم توليدها باستخدام ANTLR عند استخدامها تعطينا Parse Tree، لكننا نحتاج إلى تحويلها إلى AST خاصة بنا لنتمكّن من استخدامها في المراحل التالية من المترجم. "
        "وبموجب المتطلبات يجب بناء شجرتين مجردتين: الشجرة الأولى لـ Python، والشجرة الثانية لـ Jinja2/HTML، مع ضرورة أن يقوم المولّد بتمرير البيانات من مصفوفة البيانات في كود Python إلى الشجرة الثانية.",
    )
    add_ar(
        doc,
        "في البداية قمنا بإنشاء حزمتين: الحزمة pyast لشجرة Python، والحزمة AST لشجرة القوالب. "
        "كلتاهما مبنية وفق مفاهيم البرمجة غرضية التوجه: وراثة من عقدة جذر مجرّدة، وتعدد أشكال عبر دوال طباعة مجرّدة تنفّذها كل عقدة. "
        "وتخزّن كل عقدة اسمها المنطقي ورقم السطر بشكل صحيح كما طُلب في إعلان المشروع.",
    )
    add_ar(
        doc,
        "فالشجرة المجرّدة تحتوي قواعد اللغة ذات القيمة أو التي تحمل معنى منطقياً فقط، لا كل رموز الـ Tokens. "
        "أفضل طريقة لتمثيل كل قاعدة هي إنشاء صنف class خاص بها باعتبار القاعدة هي object (كائن) وباعتبار الرمز token هو سلسلة String، "
        "وباعتبار سلسلة القواعد سلسلة من الكائنات list of objects. وبذلك تحوي الـ AST على أصناف، وفي كل صنف ما يعبّر عن القاعدة المعبر عنها من رموز وقيم ومؤشرات على كائنات من أصناف أخرى من الـ AST، "
        "وكذلك بانٍ constructor وتابع طباعة يعيد القاعدة ومحتوياتها على شكل String وذلك لطباعة الشجرة بشكل هرمي مفهوم في النهاية.",
    )
    add_ar(
        doc,
        "في شجرة Python يكون الجذر من النوع Program، ويتفرع إلى استيرادات وتوابع مزخرفة وتعريفات توابع وجمل. "
        "ومن أصناف العقد: FunctionDef و DecoratedFunction و Assign و If و For و While و Return و Break و CallExpr و DictExpr و ListExpr وغيرها. "
        "يبني الزائر ASTBuilder هذه الشجرة انطلاقاً من Parse Tree، ويطبعها ASTPrinter بشكل مقروء.",
    )
    add_ar(
        doc,
        "أما شجرة القوالب فيكون جذرها ASTTemplate، وتحتوي محتوىً متنوّعاً: HtmlElementNode و TextNode و ASTJinjaExpression و ASTJinjaStatementNode و ASTBlock و ASTStyle وعقد CSS. "
        "يبنيها الزائر ASTVisitor الذي يرث من الصنف المولَّد htmlParserBaseVisitor. "
        "وبهذا أصبح لدينا هرمية معيّنة من الأصناف بحيث يمكننا بذلك الاستفادة من مفاهيم OOP كالـ Polymorphism والـ Abstraction وتخزين مجموعة من القواعد المختلفة مثلاً في حاوية Container واحدة.",
    )

    # ── رابعاً ───────────────────────────────────────────────────────────
    add_heading_ar(doc, "رابعاً: توليد جدول الرموز Symbol Table")

    add_ar(
        doc,
        "سوف نعتبر جدول الرموز هو عبارة عن كائن من صنف SymbolTable. وبما أن للمشروع جبهتين فقد بنينا جدولين متناسقين مع طبيعة كل لغة.",
    )
    add_ar(
        doc,
        "جدول رموز Python يحتوي على نطاق عام Global Scope ومكدّس نطاقات StackOfScopes. كل رمز (Item) هو أيضاً عبارة عن كائن مؤلَّف من معرّف واسم ونوع ومؤشر على الـ Scope الذي تم تعريفه فيه. "
        "وكل Scope هو أيضاً عبارة عن كائن مؤلَّف من معرّف واسم ونوع ومؤشر على الـ Scope الأب، ومؤشر على الرموز التي أُدخلت للتنفيذ إلى هذا الـ Scope. "
        "وكائن الـ Symbol Table إضافة إلى القائمة والمكدس السابق ذكرهما يحتوي على دوال لإضافة عناصر والتحقق من وجود عنصر في نطاق ما أو قابلية الوصول إليه من ذلك النطاق وغير ذلك من التوابع اللازمة لعمل الـ SymbolTable.",
    )
    add_ar(
        doc,
        "ثم في تابع الزائر عند تعريف أي متغير أو تابع أو رمز ما يتمّت إضافته إلى كائن SymbolTable الخاص بالبرنامج. "
        "يُهيَّأ النطاق العام مسبقاً برموز Flask (مثل request و render_template و url_for و Flask) وبدوال Python المدمجة (مثل float و int و str و len و max)، حتى لا تُعدّ هذه الأسماء «غير معرّفة» أثناء التحليل الدلالي.",
    )
    add_ar(
        doc,
        "أما جدول رموز القوالب فيخزّن رموز مستوى القالب فقط: متغيرات Jinja، والتوابع المسجّلة مع تواقيعها، وأسماء المعاملات، والكتل block. "
        "عناصر HTML وخصائصها لا تُخزَّن هنا لأنها تعيش في الـ AST. يتكوّن الجدول من شجرة نطاقات: النطاق الجذر للقالب، ونطاقات أبناء لكل {% block %}. "
        "قبل المشي على الشجرة تُسجَّل الدوال المدمجة مثل url_for عبر FunctionRegistry، ثم يملأ HtmlSymbolTableVisitor الرموز المشاهدة في القالب.",
    )
    add_ar(
        doc,
        "في نهاية الترجمة يُصدَّر الجدولان إلى ملفين نصيين مقروءين: compiler_output/symbol_table_python.txt و compiler_output/symbol_table_jinja.txt، "
        "إضافة إلى طباعة الشجرة الكاملة في الطرفية، استجابة لمتطلب «كتابة توابع لطباعة كل عقدة وبنائها بشكل مقروء، وتابع لطباعة الشجرة الكاملة مع جدول الرموز».",
    )

    # ── خامساً ───────────────────────────────────────────────────────────
    add_heading_ar(doc, "خامساً: معالجة الأخطاء Error Handling")

    add_ar(
        doc,
        "احتجنا إلى طبقات منفصلة إضافة إلى جدول الرموز لتحقيق معالجة الأخطاء. فالمتطلب ينص على وجوب معالجة الأخطاء الدلالية في كلا الجزأين (معالجة خمسة أخطاء دلالية على الأقل).",
    )
    add_ar(
        doc,
        "الصنف الأول يرث من صنف BaseErrorListener الذي تم توليده من قبل ANTLR بحيث تمرّ عبره كل الأخطاء اللفظية والنحوية ويتمّت إضافتها إلى مجمع موحّد. "
        "في مسار Python يُستخدم PyErrorListener لهذا الغرض، ويصنّف CompilerError الأخطاء إلى لفظية ونحوية ودلالية. "
        "إذا وُجدت أخطاء لفظية أو نحوية يتوقف المترجم قبل بناء الـ AST لأن الشجرة الناتجة تكون غير موثوقة، ويُكتب تقرير منظم بدل أن يفشل البرنامج بصمت.",
    )
    add_ar(
        doc,
        "أما التحليل الدلالي فيعتمد على زائر يمشي على الـ AST الخاصة بنا (وليس على Parse Tree)، ويدير النطاقات، ويستدعي صنفاً مستقلاً لكل نوع فحص. "
        "أداة ANTLR تولّد معالجة أخطاء أساسية بناءً على القواعد التي كتبناها، لكن يوجد الكثير من الأخطاء الإضافية التي نحتاج إلى معالجتها في اللغة والتي قد يكون من الصعب جداً التعبير عنها باستخدام القواعد وحدها، "
        "وقد تؤدي إلى إضافة عدد كبير من القواعد وتعديل كبير في بنية مترجمنا لمجرد تحقيق التقاط نوع واحد من الأخطاء؛ الأمر الذي يمكن تحقيقه بعدة أسطر بسيطة من الكود خلال عملية زيارة العقد وترجمتها نفسها. "
        "وهذه هي الفائدة الأكبر من تحقيق معالجة الأخطاء Error Handling بهذه الطريقة حيث تعطي مرونة عالية للغة وسهولة في التقاط ومعالجة الحالات الخاصة والـ Edge Cases.",
    )

    add_ar(doc, "وندوّن هنا حالات الأخطاء التي تمّت معالجتها في جزء Python:", size=14, bold=True, align="right", space_before=8)
    add_numbered(doc, 1, "عند استخدام متغير غير معرّف.")
    add_numbered(doc, 2, "عند تعريف تابع بنفس الاسم في النطاق نفسه (إعادة تعريف دالة).")
    add_numbered(doc, 3, "عند وجود أكثر من معامل للتابع بنفس الاسم.")
    add_numbered(doc, 4, "عند الإسناد المركّب (مثل +=) لمتغير لم يتم تعريفه بعد.")
    add_numbered(doc, 5, "عند استدعاء تابع غير معرّف من قبل.")
    add_numbered(doc, 6, "عند استدعاء متغير كأنه تابع.")
    add_numbered(doc, 7, "عند ورود break خارج جسم حلقة.")
    add_numbered(doc, 8, "عند إسناد قيمة فوق اسم تابع معرَّف فعلاً.")
    add_numbered(doc, 9, "عند ورود return خارج جسم تابع.")
    add_numbered(doc, 10, "عند تغطية اسم مستورد بتعريف متغير لاحق.")
    add_numbered(doc, 11, "عند تعارض الأنواع في تعبير ثنائي.")

    add_ar(doc, "وندوّن هنا حالات الأخطاء التي تمّت معالجتها في جزء القوالب Jinja/HTML:", size=14, bold=True, align="right", space_before=8)
    add_numbered(doc, 1, "عند استخدام متغير غير معرّف داخل {{ }}.")
    add_numbered(doc, 2, "عند استدعاء تابع غير معرّف.")
    add_numbered(doc, 3, "عند نقص وسيط إلزامي موضعي أو مسمّى في استدعاء تابع.")
    add_numbered(doc, 4, "عند تمرير وسيط مسمّى غير متوقع لتوقيع التابع.")
    add_numbered(doc, 5, "عند استدعاء تابع غير آمن داخل قيمة خاصية HTML.")
    add_numbered(doc, 6, "عند تعريف كتلة block بنفس الاسم أكثر من مرة في النطاق نفسه.")
    add_numbered(doc, 7, "عند عدم استخدام متغير الحلقة داخل جسم {% for %}.")
    add_numbered(doc, 8, "عند تكرار قيمة الخاصية id على أكثر من عنصر HTML.")
    add_numbered(doc, 9, "عند تداخل وسوم HTML بشكل غير صالح (مثل <a> داخل <a> أو <li> خارج قائمة).")

    add_ar(
        doc,
        "وقد أضفنا المستمع الخاص بنا إلى الـ Parser الذي نقوم بإنشائه في بداية البرنامج قبل بدء عملية الترجمة للتأكد من مرور جميع الأخطاء عبره خلال عملية الترجمة. "
        "كذلك مرّرنا غرض معالج الأخطاء إلى الزائر الأساسي الذي بدوره سيُمرَّر بشكل دوري لكل الزوّار الذي سيُحتاج إليهم خلال عملية الترجمة، "
        "والهدف من ذلك إمكانية إضافة الأخطاء الخاصة بنا حال وجودها في الكود الذي تتم ترجمته. "
        "وفي النهاية تُكتب كل الأخطاء في compiler_output/semantic_report.txt مصنّفة: لفظي/نحوي لـ Python، ثم دلالي لـ Python، ثم دلالي لقالب Jinja.",
    )

    # ── سادساً ───────────────────────────────────────────────────────────
    add_heading_ar(doc, "سادساً: توليد الكود Code Generation")

    add_ar(
        doc,
        "يتم توليد الكود بلغة الهدف عن طريق مسار مزدوج يضمن أن الأجزاء المولَّدة قادرة على العمل معاً، كما نصّ إعلان المشروع. "
        "المرحلة الأولى تستخرج البيانات من شجرة Python، والمرحلة الثانية تمرّر هذه البيانات إلى شجرة Jinja ثم تنتج HTML نهائياً، "
        "مع ملفات تشغيل Flask وملفات دعم الواجهة.",
    )
    add_ar(
        doc,
        "القالب الأساسي هنا ليس ملف HTML واحداً يُحقن فيه إطار جاهز، بل مجموعة قوالب إدخال حقيقية templates/*.jinja تُحلَّل عبر htmlLexer و htmlParser، "
        "ثم تُبنى AST عبر ASTVisitor، ثم يُفحصها SemanticAnalyzer، ثم ينفّذها JinjaRenderer مقابل بيانات السياق. "
        "هذا يضمن أن التوليد ليس مجرد نصوص HTML ثابتة مكتوبة داخل Java، بل ناتج حقيقي لمراحل المترجم.",
    )

    add_ar(doc, "ربط سياق Python بشجرة Jinja:", size=15, bold=True, align="right", space_before=8)
    add_ar(
        doc,
        "دالة extractData تزور Python AST بحثاً عن الإسناد products = [...] وتحوّل كل DictExpr إلى ProductData. "
        "ثم تُحوَّل القائمة إلى سياق Map تحت المفتاحين products و product (الأخير لمعاينة صفحة التفاصيل عند التوليد الأول). "
        "وهذا السياق يُمرَّر مباشرة إلى JinjaRenderer.render. وبذلك يحقّق المولّد شرط «تمرير البيانات من مصفوفة البيانات في كود Python إلى الشجرة الثانية».",
    )

    add_ar(doc, "قواعد التقابل والأمثلة:", size=15, bold=True, align="right", space_before=8)
    add_ar(
        doc,
        "إضافة لما سبق، فالكود الذي يتم توليده من قواعد التقابل والقوالب الصغيرة الناتجة عنها يتم توليده بناءً على كود الدخل بشكل دوري لضمان صحة منطق تسلسل الكود الناتج. "
        "وفيما يلي نذكر بعض الأمثلة على سبيل الاستئناس والفهم وليس الحصر، تجنّباً للإطالة عن قواعد التقابل التي قمنا باستنتاجها وتحقيقها:",
    )

    add_ar(doc, "❖ الكود:", size=14, bold=True, align="right")
    add_code(doc, 'products = [{"id": 1, "name": "Laptop", "price": 1200}, ...]')
    add_ar(
        doc,
        "الذي يُفسَّر بقاعدة الإسناد وقائمة القواميس من ملف pyParser.g4 يولّد قالب السياق:",
    )
    add_code(doc, "context[\"products\"] = [ {id, name, price, ...}, ... ]")
    add_ar(
        doc,
        "حيث سيتم استبدال قائمة المنتجات في السياق بالقيم المستخرجة من شجرة Python، وهي في هذه الحالة كائنات المنتج الفعلية وليس نصاً ثابتاً.",
    )

    add_ar(doc, "❖ الكود:", size=14, bold=True, align="right", space_before=8)
    add_code(doc, "{% for product in products %}\n  <td>{{ product.name }}</td>\n{% endfor %}")
    add_ar(
        doc,
        "الذي يُفسَّر بقاعدة for_stmt من ملف htmlParser.g4 يولّد القالب:",
    )
    add_code(doc, "for item in context[iterable]:\n    loopCtx[iterVar] = item\n    render(body, loopCtx)")
    add_ar(
        doc,
        "حيث سيتم استبدال متغير التكرار باسم المنتج في هذه الحالة product، واستبدال القائمة القابلة للتكرار بـ products، "
        "واستبدال جسم الحلقة بسلسلة العقد الواقعة بين for و endfor. وفي كل دورة تُقيَّم {{ product.name }} من الحقل name داخل عنصر السياق الحالي.",
    )

    add_ar(doc, "❖ الكود:", size=14, bold=True, align="right", space_before=8)
    add_code(doc, "{{ url_for('product_detail', product_id=product.id) }}")
    add_ar(
        doc,
        "الذي يُفسَّر بقاعدة function_call داخل jinja_expression يولّد القالب:",
    )
    add_code(doc, "/products/{product_id}")
    add_ar(
        doc,
        "حيث سيتم استبدال اسم نقطة النهاية product_detail بمسار حقيقي من جدول المسارات الثابت المطابق لـ app.py المولَّد، "
        "ثم استبدال {product_id} بقيمة الحقل id للمنتج الحالي. أما url_for('static', filename='style.css') فيُحوَّل إلى /static/style.css. "
        "والتابع url_for الذي يتم استدعاؤه في القالب السابق هو تابع معرّف مسبقاً في سجل الدوال المدمجة للغة.",
    )

    add_ar(doc, "❖ الكود:", size=14, bold=True, align="right", space_before=8)
    add_code(doc, '<img src="{{ product.image }}" alt="{{ product.name }}" />')
    add_ar(
        doc,
        "الذي يُفسَّر بقاعدة العنصر HTML (ومنها العناصر الفارغة void مثل img) يولّد عنصراً نهائياً بعد تقييم خصائص Jinja. "
        "وعلى نفس المشكلة تتم معالجة كل وسوم HTML/Jinja التي تدعمها لغتنا حيث الوسوم المتداخلة والبنى المعقّدة تتم معالجتها باستدعاءات متداخلة لتوابع مشابهة لهذا التابع المعروض، "
        "بحيث لكل وسم تدعمه اللغة تابع خاص به، وينتج عن هذه الاستدعاءات في النهاية بنية متداخلة من عناصر الواجهة يتم عرضها في المتصفح تبعاً لما أراده مستخدم كاتب الكود.",
    )

    add_ar(doc, "❖ الكود:", size=14, bold=True, align="right", space_before=8)
    add_code(doc, "@app.route(\"/products/add\", methods=[\"GET\", \"POST\"])\ndef add_product(): ...")
    add_ar(
        doc,
        "الذي يُفسَّر بقواعد المزخرف وتعريف التابع من ملف pyParser.g4 لا يُترجم إلى HTML مباشرة، بل يُستخدم لفهم بنية التطبيق. "
        "أما التطبيق التفاعلي نفسه فيُولَّد في app.py عبر GeneratedFlaskApp بحيث تُقرأ المنتجات من data/products.json، "
        "ومسارات الإضافة والحذف تستدعي خادم Java لإعادة التوليد.",
    )

    add_ar(doc, "خادم إعادة التوليد والواجهات:", size=15, bold=True, align="right", space_before=10)
    add_ar(
        doc,
        "بعد أول توليد يصبح مصدر الحقيقة للبيانات هو output/data/products.json وليس ملف Python الأصلي. "
        "الصنف RegenServer هو خادم HTTP مبني على HttpServer المدمج في JDK دون مكتبات خارجية. "
        "يستمع على المنفذ 8090 إلى POST /products/add و POST /products/delete، يحدّث JSON، ثم يعيد استخدام المسار نفسه "
        "Jinja Parser ثم AST ثم JinjaRenderer لإعادة توليد output/index.html و output/product_detail.html، "
        "ويُسجَّل سطر في compiler_output/generation_log.txt بسبب العملية ونتيجتها.",
    )
    add_ar(
        doc,
        "أما الواجهات والتنقل فقد صُمّمت لتكون متكاملة كما طُلب: عرض المنتجات، إضافة منتج، عرض تفاصيل منتج، وحذف منتج، "
        "مع روابط تنقل مشتركة بين الصفحات. تطبيق Flask الحيّ يعمل على المنفذ 5000 ويستخدم قوالب تشغيل في output/templates، "
        "بينما ملفات output/*.html هي الناتج الصريح لمرحلة التوليد من المترجم.",
    )

    # ── سابعاً ───────────────────────────────────────────────────────────
    add_heading_ar(doc, "سابعاً: استخدام كل ما سبق في البرنامج النهائي")

    add_ar(
        doc,
        "في النهاية قمنا بتعريف صنف Main وتابع main حيث يتم فيه تحديد مسار ملفات البرنامج المراد عمل compile لها. "
        "ومن أجل كل ملف من الملفات المحددة يتم قراءة الملف المطلوب وتمريره في الـ Lexer والـ Parser مع استخدام الـ SymbolTable والـ ErrorHandler بالشكل الصحيح. "
        "فيتم في النهاية توليد الـ AST الخاصة بنا وجدول الرموز SymbolTable للبرنامج المدخل مع معالجة معظم الأخطاء النحوية والدلالية، "
        "حيث يتم إظهار كل الأخطاء النحوية للمستخدم إن وُجدت، وإن لم توجد يتم إظهار كل الأخطاء الدلالية، وإن لم توجد كذلك يتم إخراج الـ AST بشكل مفهوم يعبّر عن شكل البرنامج الفعلي "
        "في ملفين باسم ast_python.json و ast_jinja.json داخل compiler_output، "
        "وكذلك إخراج الـ SymbolTable الذي يحوي على كل الرموز التي يتم تعريفها في بداية البرنامج من قبل المترجم الخاص بنا كـ Globals من أصل اللغة إضافة إلى كل الرموز التي تم تعريفها عبر البرنامج المدخل "
        "وذلك أيضاً في ملفين باسم symbol_table_python.txt و symbol_table_jinja.txt، "
        "كما وأن الكود الناتج عن عملية الترجمة بلغة الهدف يتم توليده وإخراجه أيضاً في مجلد output: صفحات HTML المسطّحة، وتطبيق Flask، وملفات CSS/JS، وبيانات المنتجات.",
    )
    add_ar(
        doc,
        "صيغة التشغيل تجمع المرحلتين تحت مدخل واحد: اللغة ثم الملف ثم اختيارياً مجلد الخرج ومجلد القوالب. "
        "عند اختيار python يمر PythonRunner بالمراحل بالترتيب: تحليل لفظي ونحوي، بناء AST، طباعة الشجرة، تحليل دلالي وتصدير جدول الرموز، ثم توليد الكود. "
        "وعند اختيار html يمكن فحص قالب منفرد عبر HtmlRunner.",
    )

    add_code(
        doc,
        "java -cp \"build/classes;dependencies/antlr-4.13.2-complete.jar\" Main python app.py\n"
        "java -cp \"...\" server.RegenServer 8090 output templates\n"
        "python output/app.py",
    )

    add_ar(
        doc,
        "حيث يمكن فتح الناتج في المتصفح مباشرة لرؤية البرنامج الناتج والتفاعل معه، أو رفعه على سيرفر ليتمكّن الناس من الوصول إليه عبر شبكة الإنترنت. "
        "التطبيق الناتج يعرض المنتجات في جدول مع صورة واسم وسعر وروابط عرض التفاصيل والحذف، وصفحة إضافة بمنوذج، وصفحة تفاصيل، مع تنقل بين هذه الواجهات. "
        "وعند إضافة منتج أو حذفه من المتصفح يستدعي Flask خادم Java تلقائياً فتتحدّث صفحات الخرج فوراً.",
    )

    add_ar(doc, "مخطط مسار الترجمة باختصار:", size=15, bold=True, align="right", space_before=8)
    add_code(
        doc,
        "app.py  →  pyLexer / pyParser  →  Python AST  →  Symbol Table + Semantic\n"
        "        →  extract products[]\n"
        "templates/*.jinja  →  htmlLexer / htmlParser  →  Jinja AST  →  Semantic\n"
        "        →  JinjaRenderer(context from Python)  →  output/*.html\n"
        "Add/Delete  →  RegenServer  →  products.json  →  re-render templates",
    )

    add_ar(
        doc,
        "وبهذا يكتمل تحقيق مراحل المترجم الأساسية المطلوبة في الإعلان: بناء القواعد، بناء الشجرتين، هيكلية العقد غرضية التوجه، "
        "التحليل الدلالي في الجزأين، توليد الكود بحيث تعمل الأجزاء معاً، تصميم واجهات وتنقّل متكامل، وطباعة العقد والشجرة وجدول الرموز.",
    )

    add_ar(doc, "النهاية", size=18, bold=True, align="center", space_before=28, space_after=6)

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
